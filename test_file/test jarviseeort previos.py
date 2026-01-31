import os
import asyncio
import logging
from docx import Document
from livekit.agents import function_tool
from playwright.async_api import async_playwright

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Ensure Reports directory exists
DATA_DIR = os.path.join(os.path.dirname(__file__), "Data")
REPORTS_DIR = os.path.join(DATA_DIR, "Reports")
if not os.path.exists(REPORTS_DIR):
    os.makedirs(REPORTS_DIR)

# DeepSeek Credentials (Hardcoded for this session as requested)
DS_EMAIL = "solisubhash2025@gmail.com"
DS_PASS = "Subhash123@"

@function_tool
async def generate_deepseek_report(topic: str) -> str:
    """
    Generates a detailed, professional report or article on ANY topic using DeepSeek's AI.
    Use this tool whenever the user asks for a "report", "document", "deep research", or "long-form writing".
    This tool opens a browser, performs the research/writing, and saves it as a Word (.docx) file.
    """
    logger.info(f"Generating report on: {topic}")
    
    async with async_playwright() as p:
        # Launch browser (Headless=False so user can see what's happening, or True for background)
        # Using False for now so user sees the "magic"
        browser = await p.chromium.launch(headless=False)
        context = await browser.new_context()
        page = await context.new_page()
        
        try:
            # 1. Login
            logger.info("Navigating to login...")
            await page.goto("https://chat.deepseek.com/login")
            
            # Wait for email input
            # Inspection says type is 'text' or placeholder 'Phone number / email address'
            # We use a robust selector
            logger.info("Waiting for email input...")
            await page.wait_for_selector('input[type="text"]', timeout=15000)
            await page.fill('input[type="text"]', DS_EMAIL)
            
            # Password
            await page.fill('input[type="password"]', DS_PASS)
            
            # Click Login Button
            # Selector: button with text "Log in"
            await page.click('button:has-text("Log in")')
            
            # Wait for chat interface (look for textarea)
            logger.info("Waiting for chat interface...")
            await page.wait_for_selector('textarea', timeout=30000)
            
            # --- Enable DeepThink (R1) and Search (Web) ---
            logger.info("Enabling DeepThink and Search...")
            
            # Define Button Locators
            # We use text filtering to be robust against order changes
            btn_deepthink = page.locator('.ds-toggle-button').filter(has_text="DeepThink")
            btn_search = page.locator('.ds-toggle-button').filter(has_text="Search")
            
            # Helper to toggle on
            async def ensure_active(btn, name):
                try:
                    if await btn.count() > 0:
                        # Check if active
                        classes = await btn.get_attribute('class')
                        if 'ds-toggle-button--selected' not in classes:
                            logger.info(f"Activating {name}...")
                            await btn.click()
                            # Wait for it to become active
                            try:
                                await btn.wait_for(state="visible") # Ensure generic visibility
                                # Optional: verify class change? 
                                # DeepSeek might take a ms to update class.
                                await asyncio.sleep(0.5)
                            except:
                                pass
                        else:
                            logger.info(f"{name} is already active.")
                    else:
                        logger.warning(f"{name} button not found.")
                except Exception as e:
                    logger.warning(f"Failed to toggle {name}: {e}")

            await ensure_active(btn_deepthink, "DeepThink")
            await ensure_active(btn_search, "Search")
            
            # --- Step 1: Initial Context & Draft ---
            logger.info("Sending Step 1 Prompt (Initial Draft)...")
            prompt_1 = f"I need a comprehensive report on '{topic}'. First, just outline the key sections, technical concepts, and core arguments. Do not write the full report yet, just set the strong foundation."
            await page.fill('textarea', prompt_1)
            await page.keyboard.press('Enter')
            
            # Wait for Step 1 generation
            await wait_for_generation(page, "step_1")

            # --- Step 2: Final Comprehensive Report with Formulas ---
            logger.info("Sending Step 2 Prompt (Deep Refinement)...")
            # Step 2: Refinement (The "Deep" part)
            # We ask for a "Markdown Code Block" to get raw tables and formulas.
            # CRITICAL: We explicitly forbid "vertical lists" for variables to prevent the layout issue.
            refinement_prompt = (
                "Great. Now, write the **FINAL COMPREHENSIVE REPORT** based on that draft. "
                "1.  Expand significantly (aim for 2500+ words). "
                "2.  Include a **Detailed Table** comparing key concepts or data. "
                "3.  Use **Mathematical Formulations** (e.g., $$ E=mc^2 $$). "
                "4.  **VARIABLE DEFINITIONS**: When defining variables (e.g. 'where P is pressure'), keep them on ONE LINE or use a Table. Do NOT list them as 'P' (newline) '=' (newline) 'Pressure'. "
                "5.  Add a 'References' section. "
                "6.  **CRITICAL OUTPUT FORMAT**: Provide the entire report inside a SINGLE MARKDOWN CODE BLOCK (```markdown ... ```). "
            )
            
            # Focus textarea again if needed (DeepSeek usually keeps focus, but safety first)
            await page.click('textarea') 
            await page.fill('textarea', refinement_prompt)
            await page.keyboard.press('Enter')
            
            # Wait for Step 2 generation (The Final Result)
            await wait_for_generation(page, "step_2")

            # Scrape content (Priority: RAW Code Blocks via <pre>)
            report_content = ""
            
            # 1. Try to find the Markdown Code Block (in a <pre> tag)
            pre_elements = await page.query_selector_all('pre')
            if pre_elements:
                logger.info("Found pre blocks (Code). Scraping raw markdown...")
                try:
                    # Try to get code element specifically to avoid headers like "Copy Download"
                    code_el = await pre_elements[-1].query_selector('code')
                    if code_el:
                        report_content = await code_el.inner_text()
                    else:
                        report_content = await pre_elements[-1].inner_text()
                except Exception as e:
                    logger.warning(f"Failed to access pre block: {e}")
                    report_content = ""
            
            # Retry scraping loop to handle stale elements
            for attempt in range(5):
                try:
                    # Refresh elements list on every attempt
                    elements = await page.query_selector_all('.ds-markdown')
                    if not elements:
                        elements = await page.query_selector_all('div[class*="markdown"]')
                    
                    if elements:
                        max_len = 0
                        for el in elements:
                            try:
                                text = await el.inner_text()
                                if len(text) > max_len:
                                    max_len = len(text)
                                    report_content = text
                            except:
                                continue # Single element fail, try others
                        
                        if len(report_content) > 200:
                            logger.info(f"Scraped longest block: {max_len} chars")
                            break # Success
                        
                    await asyncio.sleep(1.0) # Wait before retry
                except Exception as e:
                    logger.warning(f"Scraping attempt {attempt} failed: {e}")
                    await asyncio.sleep(1.0)

            # Fallback: Content is still too short (maybe failed?)
            if not report_content or len(report_content) < 200:
                 logger.warning("Main scraping failed or too short. Trying catch-all strategy.")
                 frames = await page.query_selector_all('.ds-chat-item') 
                 if frames:
                     # Check the last few frames
                     for frame in reversed(frames):
                         text = await frame.inner_text()
                         if len(text) > 200:
                             report_content = text
                             break
                     if not report_content:
                        report_content = await frames[-1].inner_text()
                 else:
                     report_content = await page.inner_text('body')

            if not report_content or len(report_content) < 100:
                logger.error(f"Scraping failed completely. Content: {report_content[:100]}")
                return f"❌ Failed to scrape report. The AI generated text, but I couldn't copy it. Found {len(report_content)} chars."
            
            # 5. Save to Docx with Formatting
            doc_path = os.path.join(REPORTS_DIR, f"{topic.replace(' ', '_')}_Report.docx")
            doc_path = os.path.abspath(doc_path) # Normalize path for Windows
            
            # DEBUG: Save exact raw content to see what we are scraping
            raw_path = os.path.join(REPORTS_DIR, f"{topic.replace(' ', '_')}_RAW.txt")
            with open(raw_path, "w", encoding="utf-8") as f:
                f.write(report_content)
            logger.info(f"DEBUG: Saved raw content to {raw_path}")
            
            save_formatted_doc(report_content, doc_path, topic)
            
            logger.info(f"Report saved to {doc_path} ({len(report_content)} chars)")
            
            # Post-save buffer to allow OS file system updates
            await asyncio.sleep(1.0) 
            
            # 6. Open File (Verify existence first with retry)
            for _ in range(3):
                if os.path.exists(doc_path):
                    break
                await asyncio.sleep(1.0)
                
            if os.path.exists(doc_path):
                logger.info(f"Opening report: {doc_path}")
                os.startfile(doc_path)
                return f"✅ Detaled Report Generated! ({len(report_content)} chars). Formatting applied. file:///{doc_path.replace(os.sep, '/')}"
            else:
                logger.error(f"File was saved but not found at: {doc_path}")
                return f"✅ Report generated but could not be opened. Path: {doc_path}"

        except Exception as e:
            logger.error(f"DeepSeek automation failed: {e}")
            try:
                await page.screenshot(path=os.path.join(DATA_DIR, "error_screenshot.png"))
            except:
                pass
            return f"❌ Report generation failed: {e}"
            
        finally:
            await browser.close()

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def save_formatted_doc(text, path, topic):
    """Parses Markdown-like text and saves a formatted Word doc with advanced styling."""
    doc = Document()
    
    # Clean up "Code Block" markers if the whole thing is wrapped
    # The prompt usually results in ```markdown ... ```
    text = text.replace("```markdown", "").replace("```", "").strip()
    
    # --- Artifact Cleaning ---
    lines = text.split('\n')
    cleaned_lines = []
    
    # Words to ignore at the start of the file (UI artifacts)
    ignore_words = ["markdown", "copy", "download", "copy code", "python", "html"]
    
    start_collecting = False
    for line in lines:
        stripped_line = line.strip().lower()
        
        # If we haven't started collecting, check for artifacts
        if not start_collecting:
            if not stripped_line:
                continue
            if stripped_line in ignore_words:
                continue
            # Also ignore if the line is exactly the topic (case-insensitive)
            if stripped_line == topic.strip().lower():
                continue
            # If we passed the filters, this is real content
            start_collecting = True
            cleaned_lines.append(line)
        else:
            cleaned_lines.append(line)
            
    lines = cleaned_lines
    # Reassemble for processing (the rest of the function iterates 'lines')
    # But wait, original code does text.split('\n') later. We should just update 'text' or 'lines'.
    # The original function splits 'text' at line 247. Let's do the split logic here and use it.
    
    # Note: The original code iterates 'lines' at line 265 (merged_lines logic starts there).
    # We will let the "Line Processing" loop handle it, but we pre-process 'lines' here.
    # However, to be minimally invasive, we'll just update 'text' if possible, or intercept 'lines' later.
    # The function below does: lines = text.split('\n') at line 247.
    # So we can just update 'text' by joining 'cleaned_lines'.
    text = "\n".join(cleaned_lines)

    # --- Style Setup ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Title
    title = doc.add_heading(topic, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0) # Black
    
    # Line Processing
    lines = text.split('\n')
    
    # State flags
    in_code_block = False
    in_references = False
    table_buffer = [] # Store rows of a table being built
    
    for line in lines:
        line = line.strip()
        if not line:
            # If we were building a table and hit a blank line, verify if table is done
            if table_buffer:
                process_table(doc, table_buffer)
                table_buffer = []
            continue
        
        # --- Citation Cleanup ---
        # Fix artifacts like "- 3 -" or "- 2 - 5" which are poorly formatted refs
        # Regex: match hyphen-number-hyphen sequence
        line = re.sub(r'-\s*(\d+)\s*-', r'[\1]', line) # "- 3 -" -> "[3]"
        line = re.sub(r'-\s*(\d+)\s*-\s*(\d+)', r'[\1-\2]', line) # "- 1 - 5" -> "[1-5]"

        # --- Table Detection ---
        if line.startswith('|') and line.endswith('|'):
            table_buffer.append(line)
            continue
        else:
            # If we have a buffer but this line isn't a table row, process the table first
            if table_buffer:
                process_table(doc, table_buffer)
                table_buffer = []

        # --- Formula Detection ---
        if line.startswith('$$') or line.startswith('\\['):
            content = line.replace('$$', '').replace('\\[', '').replace('\\]', '').strip()
            if content:
                p = doc.add_paragraph(content)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'Cambria Math'
                    run.font.size = Pt(12)
                    run.italic = True
            continue

        # --- Headlines ---
        if line.startswith('#'):
            level = line.count('#')
            content = line.strip('# ').strip()
            level = min(level, 3) 
            h = doc.add_heading(content, level=level)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue
            
        # --- List Items ---
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return_to_normal_style(p)
            continue
            
        if line[0].isdigit() and line[1:3] in ['. ', ') ']:
            p = doc.add_paragraph(line, style='List Number')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return_to_normal_style(p)
            continue
            
        # --- References Highlighting ---
        if "references" in line.lower() and len(line) < 20: 
             # Likely a header "References" or "## References"
             # Just explicitly bold it if it wasn't a markdown header
             pass

        # --- Normal Paragraph ---
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Bold Parsing
        parts = line.split("**")
        for i, part in enumerate(parts):
            if not part: continue
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            if i % 2 == 1:
                run.bold = True
    
    # Flushing any remaining table
    if table_buffer:
        process_table(doc, table_buffer)

    # Save
    doc.save(path)

def process_table(doc, lines):
    """Converts a list of markdown table lines into a Word table."""
    # Remove alignment row (e.g. |---|---|)
    valid_rows = [row for row in lines if '---' not in row]
    
    if not valid_rows:
        return

    # Determine dimensions
    # Split by | and filter empty strings (since |row| splits to ['', 'row', ''])
    headers = [c.strip() for c in valid_rows[0].split('|') if c.strip()]
    cols = len(headers)
    rows = len(valid_rows)
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid' # Standard borders
    
    for r_idx, row_str in enumerate(valid_rows):
        cells = [c.strip() for c in row_str.split('|') if c.strip()]
        # Safety check for column count mismatch
        for c_idx, text in enumerate(cells):
            if c_idx < cols:
                cell = table.cell(r_idx, c_idx)
                # cell.text = text # Basic
                # Formatting cell text
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                if r_idx == 0: # Header
                     run.bold = True
                     cell.vertical_alignment = 1 # Center

def return_to_normal_style(paragraph):
    """Ensures list items still use Times New Roman."""
    for run in paragraph.runs:
         run.font.name = 'Times New Roman'

async def wait_for_generation(page, step_name):
    """Helper to wait for generation to complete."""
    logger.info(f"Waiting for {step_name} generation...")
    
    # 1. Wait for 'Stop' button (Start)
    try:
        await asyncio.sleep(3)
    except:
        pass

    # 2. Wait for Stability
    previous_len = 0
    stable_count = 0
    
    # Increase wait time to 5 minutes (100 * 3s = 300s)
    # DeepThink + Search can take a very long time.
    for i in range(100): 
        await asyncio.sleep(3)
        
        elements = await page.query_selector_all('.ds-markdown')
        if not elements:
            elements = await page.query_selector_all('div[class*="markdown"]')
        
        current_len = 0
        if elements:
            try:
                current_text = await elements[-1].inner_text()
                current_len = len(current_text)
            except:
                pass # Element might have detached
        
        logger.info(f"[{step_name}] Check {i}: {current_len} chars")

        if current_len > 100 and current_len == previous_len:
            stable_count += 1
            # Require 15 seconds of silence (5 checks) to confirm it is actually done
            # This prevents existing early if it pauses to "Think".
            if stable_count >= 5:
                logger.info(f"[{step_name}] Generation complete (Stable).")
                break
        else:
            stable_count = 0
        
        previous_len = current_len
    
    # Screeny after each step
    await page.screenshot(path=os.path.join(DATA_DIR, f"debug_{step_name}_done.png"))
