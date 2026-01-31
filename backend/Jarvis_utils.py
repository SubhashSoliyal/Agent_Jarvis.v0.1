import os
import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from livekit.agents import function_tool
from backend.Jarvis_prompts import personal_context

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@function_tool
async def get_user_details() -> str:
    """
    Retrieves the user's personal details (Name, College, ID, Address, etc.).
    Use this tool ONLY when you need to fill in personal information for a task 
    (e.g., writing a letter, application, or email) or if the user specifically asks "What is my college ID?".
    """
    return f"Here are the user's personal details:\n{personal_context}"

def save_formatted_doc(content: str, path: str, topic: str):
    """
    Saves content to a formatted .docx file.
    Supports basic Markdown-like headers, bold text, and lists.
    """
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading(topic, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process content line by line for basic formatting
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.save(path)
        logger.info(f"Document saved successfully to {path}")
        return True
    except Exception as e:
        logger.error(f"Error saving document: {e}")
        return False

def clean_browser_user_data(user_data_dir: str):
    """
    Cleans up lock files in the user data directory to prevent launch errors.
    """
    import shutil
    import glob
    
    try:
        if not os.path.exists(user_data_dir):
            return

        # List of lock files to remove
        lock_files = [
            "SingletonLock",
            "SingletonCookie",
            "SingletonSocket",
            "lockfile"
        ]

        for filename in lock_files:
            file_path = os.path.join(user_data_dir, filename)
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    logger.info(f"Removed lock file: {file_path}")
                except Exception as e:
                    logger.warning(f"Failed to remove lock file {file_path}: {e}")

        # Also try to remove any temporary crashpad files if they exist
        # But be careful not to delete important data.
        
    except Exception as e:
        logger.warning(f"Error cleaning browser user data: {e}")
